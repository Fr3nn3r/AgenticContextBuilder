"""Generate the Holdout Evaluation Report as a professional Word document."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT = os.path.join(os.path.dirname(__file__), "HOLDOUT-EVAL-report-v2.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_kv_table(doc, pairs, col_widths=None, header_color="1F4E79"):
    """Two-column key-value table (no header row, left col is bold)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(pairs):
        kc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        kc.text = ""
        vc.text = ""
        kr = kc.paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(9)
        vr = vc.paragraphs[0].add_run(str(v))
        vr.font.size = Pt(9)
        set_cell_shading(kc, "E8EDF2")
        if i % 2 == 1:
            set_cell_shading(vc, "F9F9F9")
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Cm(w)
    return table


def para(doc, text, style=None, bold=False, size=None, space_after=Pt(6), alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if alignment is not None:
        p.alignment = alignment
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    run = p.runs[0]
    run.font.size = Pt(10)
    return p


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# -- Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# -- Header / Footer
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("CONFIDENTIAL")
hr.font.size = Pt(8)
hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("ContextBuilder \u2014 Holdout Evaluation Report v2")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ===== TITLE =====
doc.add_paragraph()
title = doc.add_heading("Holdout Evaluation Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("NSA Motor Warranty Claims \u2014 Dataset v2")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()
add_kv_table(doc, [
    ("Date", "3 February 2026"),
    ("Dataset", "nsa-motor-holdout (30 claims: 15 approved, 15 denied)"),
    ("Status", "Holdout evaluation \u2014 blind test on unseen claims"),
], col_widths=[5, 12])

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading("1. Executive Summary", level=1)

para(doc,
     "This report presents the results of a blind holdout evaluation of the "
     "ContextBuilder claims-processing pipeline on 30 previously unseen NSA "
     "motor warranty claims. The dataset was balanced (15 approved, 15 denied) "
     "and none of these claims were used during development or tuning.",
     size=Pt(10))

para(doc,
     "The pipeline achieved 76.7% decision accuracy (23 of 30 correct), with "
     "a false-reject rate of 26.7% and a false-approve rate of 20.0%. "
     "Of the 11 correctly approved claims, 3 had payout amounts within 5% "
     "of the ground-truth figure.",
     size=Pt(10))

doc.add_heading("Key Results", level=2)
add_styled_table(doc,
    ["Metric", "Value"],
    [
        ["Decision Accuracy", "76.7% (23 / 30)"],
        ["Approved Correct", "11 / 15"],
        ["Denied Correct", "12 / 15"],
        ["False Reject Rate", "26.7%"],
        ["False Approve Rate", "20.0%"],
        ["Payout Accuracy (within 5%)", "27.3% (3 / 11 approved-correct)"],
    ],
    col_widths=[8, 8])

doc.add_paragraph()

doc.add_heading("Assumptions & Scope Limitations", level=2)

para(doc,
     "The following assumptions and scope limitations apply to this evaluation. "
     "These represent capabilities that are planned but not yet integrated into "
     "the production pipeline:",
     size=Pt(10))

bullet(doc,
       "Parts-number lookup service: The pipeline assumes access to a parts "
       "database that maps OEM part numbers to system categories (e.g. "
       "electrical_system, engine, suspension). For this evaluation, the "
       "database was pre-populated with the relevant part numbers. In production, "
       "this will be replaced by a live lookup service or API integration. "
       "One claim (64868) was affected by a matching failure in this simulated "
       "lookup, resulting in an incorrect rejection.")

bullet(doc,
       "Policy payment-status service: The pipeline currently evaluates policy "
       "validity based on the dates and terms extracted from the guarantee "
       "document. It does not have access to a real-time service confirming "
       "whether premiums have been paid. One claim (64877) was denied by NSA "
       "due to a premium non-payment lapse \u2014 information that is only "
       "available through an internal policy-management system. The pipeline "
       "correctly referred this claim for human review, as it had insufficient "
       "data to make a determination.")

bullet(doc,
       "Policy renewal documents: The pipeline can only validate policy "
       "validity against documents present in the claim file. One claim "
       "(64887) was approved by NSA based on a policy renewal that was not "
       "included in the submitted documents. Integration with the policy "
       "management system would resolve this class of error.")

doc.add_page_break()

# ===== 2. ERROR BREAKDOWN =====
doc.add_heading("2. Error Classification", level=1)

para(doc,
     "Of the 7 decision errors, each falls into a distinct category with "
     "different root causes and remediation paths.",
     size=Pt(10))

add_styled_table(doc,
    ["#", "Category", "Claim(s)", "Fixable", "Priority"],
    [
        ["1", "False approve \u2014 excluded component not caught", "64827", "Yes", "HIGH"],
        ["2", "False approve \u2014 causal exclusion (root-cause denial)", "64943", "Hard", "LOW"],
        ["3", "False reject \u2014 parts-database matching failure", "64868", "Yes", "HIGH"],
        ["4", "False reject \u2014 sub-component not in explicit policy list", "65352", "Yes", "MEDIUM"],
        ["5", "False reject \u2014 missing policy renewal document", "64887", "No", "N/A"],
        ["6", "Referral \u2014 spurious VIN from unrelated document", "64846", "Partial", "MEDIUM"],
        ["7", "Referral \u2014 insufficient documents (1 doc in file)", "64877", "No", "N/A"],
    ],
    col_widths=[1, 7, 2.5, 2, 2.5])

doc.add_page_break()

# ===== 3. DECISION ERRORS \u2014 DETAILED =====
doc.add_heading("3. Decision Errors \u2014 Detailed Analysis", level=1)

# --- 3.1 FALSE APPROVALS ---
doc.add_heading("3.1 False Approvals", level=2)

# -- 64827 --
doc.add_heading("Claim 64827 \u2014 Land Rover RR Sport 4.4", level=3)

add_kv_table(doc, [
    ("Ground Truth", "DENIED \u2014 Hoses explicitly excluded from contract"),
    ("System Decision", "APPROVE \u2014 Primary repair (Bolzen/bolt) covered under engine"),
    ("Ground Truth Payout", "CHF 0"),
    ("System Payout", "CHF 187.64"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The system correctly extracted the policy\u2019s excluded components from "
     "the guarantee document. The exclusion list explicitly includes "
     "\u201cSchl\u00e4uche\u201d (hoses). The coverage analyser also correctly "
     "classified the hose (CHF 79.00) as NOT COVERED and the bolt "
     "(CHF 3.70) as COVERED.",
     size=Pt(10))

para(doc,
     "However, the screening logic evaluates only the primary covered component "
     "(the bolt at CHF 3.70) and confirms it is in a covered category. It does "
     "not check whether the main repair \u2014 the hose at CHF 79.00 \u2014 is an "
     "explicitly excluded item. This is a design gap: when the highest-value "
     "item is excluded but a minor covered part exists, the system should "
     "reject rather than approve.",
     size=Pt(10))

add_styled_table(doc,
    ["Line Item", "Part Code", "Price (CHF)", "Coverage", "Reasoning"],
    [
        ["Schlauch (hose)", "LR066536", "79.00", "NOT COVERED", "Hoses explicitly excluded"],
        ["Bolzen (bolt)", "LR033655", "3.70", "COVERED", "Engine category"],
    ],
    col_widths=[4, 3, 2.5, 3, 4.5])

doc.add_paragraph()
p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Add an excluded-component check to the decision logic. When the highest-value "
    "line item is an explicitly excluded component, the system should reject the "
    "claim regardless of whether a minor covered part exists. The excluded-components "
    "data is already extracted; it simply needs to be used in the decision path."
).font.size = Pt(10)

doc.add_paragraph()

# -- 64943 --
doc.add_heading("Claim 64943 \u2014 Mini Cooper (JCW R57)", level=3)

add_kv_table(doc, [
    ("Ground Truth", "DENIED \u2014 Malfunction caused by a non-insured part"),
    ("System Decision", "APPROVE \u2014 Turbocharger covered under turbo_supercharger"),
    ("Ground Truth Payout", "CHF 0"),
    ("System Payout", "CHF 1,929.10"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The cost estimate contains 30 line items totalling CHF 14,160. The turbocharger "
     "(CHF 2,598) IS a covered component. The system correctly identified it and "
     "approved at a 60% reimbursement rate. All screening checks passed.",
     size=Pt(10))

para(doc,
     "However, the ground-truth denial states: \u201cNo liability exists when the "
     "malfunction is caused by a non-insured part.\u201d This is a causal exclusion \u2014 "
     "the turbo failure was triggered by a failure in a non-covered component "
     "(likely the engine assembly or clutch). The invoice scope (full engine replacement "
     "+ turbo + clutch + flywheel) strongly suggests cascading damage from a root "
     "failure in the engine.",
     size=Pt(10))

add_styled_table(doc,
    ["Item Category", "Count", "Total (CHF)", "Coverage"],
    [
        ["Turbocharger", "1", "2,597.59", "COVERED"],
        ["Turbo kit", "1", "400.89", "COVERED"],
        ["Fasteners (engine)", "14", "305.24", "COVERED (ancillary)"],
        ["Engine assembly (complete)", "1", "7,982.19", "NOT COVERED"],
        ["Dual-mass flywheel", "1", "1,573.08", "NOT COVERED"],
        ["Clutch parts", "1", "620.75", "NOT COVERED"],
        ["Gaskets", "6", "516.08", "NOT COVERED (excluded)"],
        ["Other", "5", "163.64", "NOT COVERED"],
    ],
    col_widths=[5, 2, 3, 7])

doc.add_paragraph()
p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "This requires causal-chain reasoning that is beyond the current pipeline\u2019s "
    "capability. The system evaluates each component independently. Detecting that "
    "a covered component failed due to a non-covered root cause would require either "
    "an explicit causation flag from the documents (e.g. an expert report) or "
    "heuristic reasoning about invoice composition. Classified as hard to fix."
).font.size = Pt(10)

doc.add_paragraph()

# --- 3.2 FALSE REJECTIONS ---
doc.add_heading("3.2 False Rejections", level=2)

# -- 64868 --
doc.add_heading("Claim 64868 \u2014 Mercedes C 43 AMG 4Matic", level=3)

add_kv_table(doc, [
    ("Ground Truth", "APPROVED \u2014 CHF 793.70"),
    ("System Decision", "REJECT \u2014 Component coverage hard fail (auto-reject)"),
    ("System Payout", "CHF 0"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The claim involves a panoramic sunroof repair (12 line items). The key "
     "part is a sunroof motor (GETRIEBEMOTOR, part A205 906 41 04, CHF 450). "
     "This part exists in the parts database mapped to the electrical system "
     "category (roof motor) and should have been identified as covered.",
     size=Pt(10))

para(doc,
     "However, the parts-number lookup did not match this part. It fell through "
     "to the LLM-based classifier, which incorrectly determined it was not covered "
     "(confidence 0.70). Without a covered part to anchor the claim, the related "
     "labour was demoted, and a different part (DACHMECHANIK / roof mechanism, "
     "CHF 413, body category) was selected as the primary repair. Since \u201cbody\u201d "
     "is not a covered category, the screening triggered a hard fail.",
     size=Pt(10))

add_styled_table(doc,
    ["Line Item", "Code", "Type", "Price (CHF)", "Coverage", "Issue"],
    [
        ["GETRIEBEMOTOR (sunroof motor)", "A205 906 41 04", "Parts", "450.00",
         "NOT COVERED", "Part-number lookup failed"],
        ["DACHMECHANIK (roof mechanism)", "A205 780 01 75", "Parts", "413.00",
         "NOT COVERED", "Selected as primary (wrong)"],
        ["MOTOR SCHIEBEDACH (sunroof motor labour)", "77-6995-00", "Labour", "235.00",
         "NOT COVERED", "Correctly matched, then demoted"],
        ["F\u00dcHRUNGSSCHIENEN (guide rails labour)", "77-7002-00", "Labour", "658.00",
         "NOT COVERED", "No covered anchor"],
    ],
    col_widths=[4, 3, 1.5, 2, 2.5, 4])

doc.add_paragraph()
p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Debug the part-number matching failure for A205 906 41 04. The part is in the "
    "database but was not matched \u2014 likely a normalisation issue (spacing, dashes). "
    "Once production integration with a live parts-lookup service is in place, this "
    "class of error should not recur. This is the highest-impact single fix: correcting "
    "one part-number match would cascade into correct labour promotion, correct primary "
    "repair selection, and a correct approval."
).font.size = Pt(10)

doc.add_paragraph()

# -- 65352 --
doc.add_heading("Claim 65352 \u2014 BMW X3 xDrive 48V 20d", level=3)

add_kv_table(doc, [
    ("Ground Truth", "APPROVED \u2014 CHF 180.00"),
    ("System Decision", "REJECT \u2014 Component coverage hard fail"),
    ("System Payout", "CHF 0"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The claim is for a Partikelsensor (particulate/DPF sensor, CHF 211) plus "
     "replacement labour (CHF 119) and diagnostics (CHF 25). The policy\u2019s exhaust "
     "system category covers only two components: Katalysator (catalytic converter) "
     "and Lambda-Sonde (oxygen sensor). The DPF sensor is not in this list.",
     size=Pt(10))

para(doc,
     "The system correctly determined that the Partikelsensor is not explicitly listed. "
     "NSA\u2019s adjuster, however, applied a broader interpretation: the DPF sensor is "
     "a functional part of the exhaust system, and since the exhaust category is covered, "
     "the sensor should be covered by extension.",
     size=Pt(10))

p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Add Partikelsensor / DPF-Sensor to the exhaust sub-components in the policy "
    "configuration. This aligns the system\u2019s interpretation with NSA\u2019s practice of "
    "treating functional sub-components of covered categories as covered."
).font.size = Pt(10)

doc.add_paragraph()

# -- 64887 --
doc.add_heading("Claim 64887 \u2014 Land Rover Discovery 3.0 Si6 HSE", level=3)

add_kv_table(doc, [
    ("Ground Truth", "APPROVED \u2014 CHF 1,906.02"),
    ("System Decision", "REJECT \u2014 Policy validity hard fail"),
    ("System Payout", "CHF 0"),
], col_widths=[5, 12])

doc.add_paragraph()

add_styled_table(doc,
    ["Field", "Value"],
    [
        ["Guarantee Period", "09.09.2024 \u2013 08.09.2025"],
        ["Claim Date", "19.11.2025"],
        ["Days After Expiry", "71 days"],
    ],
    col_widths=[6, 11])

doc.add_paragraph()
para(doc,
     "The system correctly determined that the claim date falls 71 days after the "
     "policy expiration shown in the guarantee document. NSA approved this claim, "
     "which implies the policy was renewed \u2014 but the renewal document is not present "
     "in the claim file.",
     size=Pt(10))

p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "This is a data-completeness issue. Integration with the policy management "
    "system (to confirm current policy status and renewal dates) would eliminate "
    "this class of error. See also the policy payment-status service assumption "
    "noted in the Executive Summary."
).font.size = Pt(10)

doc.add_paragraph()

# --- 3.3 REFERRAL ERRORS ---
doc.add_heading("3.3 Referral Errors", level=2)

# -- 64846 --
doc.add_heading("Claim 64846 \u2014 Volkswagen Golf 2.0 TSI GTI Clubsport", level=3)

add_kv_table(doc, [
    ("Ground Truth", "APPROVED \u2014 CHF 321.05"),
    ("System Decision", "REFER TO HUMAN \u2014 VIN inconsistency detected"),
    ("System Payout (conditional)", "CHF 1,318.95"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The system detected two different VINs across the claim documents: one "
     "Volkswagen VIN (WVWZZZAUZHW146216, consistent across the cost estimate "
     "and guarantee) and one Audi VIN (WAUZZZF24KN016070) from what appears "
     "to be an unrelated document inadvertently included in the claim file.",
     size=Pt(10))

para(doc,
     "All other screening checks passed. The system conservatively referred the "
     "claim because VIN conflicts are a legitimate fraud indicator. This is a "
     "defensible decision \u2014 the referral is cautious rather than wrong.",
     size=Pt(10))

p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Improve VIN source attribution by weighting VINs from the guarantee and cost "
    "estimate higher than those from ancillary documents. When the primary documents "
    "are consistent and the conflicting VIN comes from a secondary source, the system "
    "should approve with a note rather than refer."
).font.size = Pt(10)

doc.add_paragraph()

# -- 64877 --
doc.add_heading("Claim 64877 \u2014 Volkswagen T6 Kombi 2.0 4M", level=3)

add_kv_table(doc, [
    ("Ground Truth", "DENIED \u2014 Non-payment premium lapse"),
    ("System Decision", "REFER TO HUMAN \u2014 Insufficient data"),
], col_widths=[5, 12])

doc.add_paragraph()
para(doc,
     "The claim file contains only one document: a vehicle registration certificate. "
     "No policy document, cost estimate, or repair invoice is present. With 9 of 10 "
     "screening checks skipped due to missing data, the system correctly referred the "
     "claim for human review.",
     size=Pt(10))

para(doc,
     "The ground-truth denial is based on a premium non-payment lapse \u2014 information "
     "that resides in NSA\u2019s internal policy-management system and cannot be inferred "
     "from claim documents alone.",
     size=Pt(10))

p = para(doc, "", size=Pt(10))
r = p.add_run("Remediation: ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Integration with a policy payment-status service would allow the system to "
    "detect premium lapses and issue an automatic denial. Without this integration, "
    "REFER TO HUMAN is the correct and defensible outcome. See the Executive Summary "
    "for the planned service assumption."
).font.size = Pt(10)

doc.add_page_break()

# ===== 4. AMOUNT MISMATCHES =====
doc.add_heading("4. Payout Amount Mismatches", level=1)

para(doc,
     "All 8 claims below were decided correctly (APPROVE = APPROVE) but the "
     "calculated payout amount differs from the ground truth by more than 5%.",
     size=Pt(10))

add_styled_table(doc,
    ["Claim", "Vehicle", "GT Amount (CHF)", "System Amount (CHF)", "Difference", "Diff %"],
    [
        ["64687", "BMW 335i xDrive", "2,566.99", "2,334.96", "\u2212232.03", "\u22129.0%"],
        ["64873", "Audi Q5 2.0 TFSI", "3,415.56", "3,948.38", "+532.82", "+15.6%"],
        ["64942", "Audi A6 Avant", "502.05", "438.71", "\u221263.34", "\u221212.6%"],
        ["65010", "BMW X3 xDrive30i", "673.79", "839.94", "+166.15", "+24.7%"],
        ["65044", "Seat Leon ST Cupra", "863.15", "251.27", "\u2212611.88", "\u221270.9%"],
        ["65316", "Mercedes GLA 45 AMG", "366.42", "336.23", "\u221230.19", "\u22128.2%"],
        ["65345", "Peugeot Partner 1.6 HDi", "958.42", "1,019.59", "+61.17", "+6.4%"],
        ["65356", "Land Rover Range Rover", "162.05", "126.90", "\u221235.15", "\u221221.7%"],
    ],
    col_widths=[1.8, 4, 2.5, 2.8, 2, 1.9])

doc.add_paragraph()

doc.add_heading("Common Root Causes", level=2)

bullet(doc,
       "Line-item selection: NSA adjusters manually strike specific line items from "
       "the cost estimate. The system does not have access to the marked-up estimate "
       "and must infer coverage from component matching rules.")

bullet(doc,
       "Labour association: Some claims have labour items that the system cannot "
       "link to covered parts, resulting in under-payment. Conversely, the system "
       "sometimes includes labour that NSA\u2019s adjuster excluded.")

bullet(doc,
       "Deductible formula: The deductible calculation base (before or after rate "
       "reduction, before or after VAT) may differ between the system\u2019s formula "
       "and NSA\u2019s internal calculation.")

bullet(doc,
       "Reimbursement rate: Age-based rate degradation rules (e.g. \u201cD\u00e8s 8 ans 70%\u201d) "
       "may be applied differently by the system and NSA.")

doc.add_page_break()

# ===== 5. CORRECTLY PROCESSED =====
doc.add_heading("5. Correctly Processed Claims", level=1)

para(doc,
     "The system correctly decided 23 of 30 claims. Of the 15 denied claims in the "
     "ground truth, 12 were correctly identified with matching denial reasons "
     "(component not covered, wear parts excluded, mileage exceedance, etc.). Of "
     "the 15 approved claims, 11 were correctly approved.",
     size=Pt(10))

add_kv_table(doc, [
    ("Correct Denials (12/15)", "64808, 64822, 64843, 64844, 64862, 64867, 64870, 64871, 64880, 64883, 64896, 64945"),
    ("Correct Approvals (11/15)", "64687, 64823, 64873, 64942, 65010, 65037, 65044, 65316, 65318, 65345, 65356"),
    ("Approvals Within 5% Tolerance (3/11)", "64823, 65037, 65318"),
], col_widths=[6, 11])

doc.add_page_break()

# ===== 6. RECOMMENDED NEXT STEPS =====
doc.add_heading("6. Recommended Next Steps", level=1)

doc.add_heading("High Priority", level=2)

p = doc.add_paragraph()
r = p.add_run("1. Excluded-component screening check. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "When the highest-value line item is an explicitly excluded component, "
    "the system should reject the claim. The excluded-components data is "
    "already extracted from the guarantee document \u2014 it needs to be "
    "incorporated into the decision logic. (Fixes claim 64827; prevents "
    "similar false approvals.)"
).font.size = Pt(10)

p = doc.add_paragraph()
r = p.add_run("2. Parts-database integration. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "The part-number lookup for A205 906 41 04 failed despite the part being "
    "present in the database. Debugging this normalisation issue is the "
    "single highest-impact fix: one correct match would cascade into correct "
    "labour promotion, correct primary-repair selection, and a correct "
    "approval. In production, a live parts-lookup service will replace the "
    "simulated database. (Fixes claim 64868.)"
).font.size = Pt(10)

doc.add_heading("Medium Priority", level=2)

p = doc.add_paragraph()
r = p.add_run("3. Exhaust sub-component coverage. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Add Partikelsensor / DPF-Sensor to the exhaust category\u2019s covered "
    "components to align with NSA\u2019s broader interpretation. (Fixes claim 65352.)"
).font.size = Pt(10)

p = doc.add_paragraph()
r = p.add_run("4. VIN source attribution. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Weight VINs from primary documents (guarantee, cost estimate) higher "
    "than those from ancillary documents. (Partially fixes claim 64846.)"
).font.size = Pt(10)

p = doc.add_paragraph()
r = p.add_run("5. Deductible formula alignment. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Investigate whether the deductible base should be calculated before or "
    "after rate reduction, and before or after VAT, to improve payout accuracy "
    "across multiple claims."
).font.size = Pt(10)

doc.add_heading("Planned Integrations", level=2)

p = doc.add_paragraph()
r = p.add_run("6. Policy payment-status service. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "A real-time service confirming premium payment status would allow the "
    "system to detect non-payment lapses and issue automatic denials. "
    "(Would fix claim 64877.)"
).font.size = Pt(10)

p = doc.add_paragraph()
r = p.add_run("7. Policy management system integration. ")
r.bold = True
r.font.size = Pt(10)
p.add_run(
    "Access to current policy status (including renewals) would resolve "
    "cases where the claim file lacks a renewal document. "
    "(Would fix claim 64887.)"
).font.size = Pt(10)

doc.add_heading("Not Fixable by the System", level=2)

para(doc,
     "Claim 64943 (causal exclusion reasoning) requires understanding failure "
     "chains across components \u2014 a capability beyond the current pipeline. This "
     "type of claim requires expert human judgement.",
     size=Pt(10))

# ===== SAVE =====
doc.save(OUTPUT)
print(f"Report saved to: {OUTPUT}")
