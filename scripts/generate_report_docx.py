"""Generate a professional Word document from the NSA claims data analysis report."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pathlib

SCRIPT_DIR = pathlib.Path(__file__).resolve().parent.parent
OUTPUT = SCRIPT_DIR / "docs" / "NSA-Claims-Data-Analysis-Report.docx"
CHARTS_DIR = SCRIPT_DIR / "docs" / "report_charts"

# -- Colors --
DARK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x00, 0x52, 0x8A)
LIGHT_ACCENT = "D6E8F5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell border. kwargs: top, bottom, left, right with dict of sz, color, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "CCCCCC")}"/>'
        )
        tcBorders.append(el)


def style_header_row(row, col_count):
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, "00528A")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None, right_align_cols=None):
    """Add a formatted table. right_align_cols is a set of 0-based column indices."""
    right_align_cols = right_align_cols or set()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = "Arial"
        set_cell_shading(cell, "00528A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F5F8FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if c_idx in right_align_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Arial"
            run.font.color.rgb = BODY_COLOR
            set_cell_shading(cell, bg)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    # Reduce cell padding
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)

    return table


def add_heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = ACCENT
        run.font.name = "Arial"
    return p


def add_heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = "Arial"
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Arial"
    run.font.color.rgb = BODY_COLOR
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_body_bold_lead(doc, bold_text, rest_text):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.font.size = Pt(10.5)
    r1.font.name = "Arial"
    r1.font.bold = True
    r1.font.color.rgb = BODY_COLOR
    r2 = p.add_run(rest_text)
    r2.font.size = Pt(10.5)
    r2.font.name = "Arial"
    r2.font.color.rgb = BODY_COLOR
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.size = Pt(10.5)
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.color.rgb = BODY_COLOR
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.name = "Arial"
        r2.font.color.rgb = BODY_COLOR
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Arial"
        run.font.color.rgb = BODY_COLOR
    p.paragraph_format.space_after = Pt(3)
    return p


def add_callout(doc, text):
    """Add an indented, italic insight paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x00, 0x52, 0x8A)
    return p


def add_chart(doc, filename, width=Inches(5.5)):
    """Add a chart image centered in the document."""
    path = CHARTS_DIR / filename
    if not path.exists():
        print(f"  WARNING: chart not found: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.add_run().add_picture(str(path), width=width)


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -- Default font --
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_COLOR
    style.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NSA Motor Guarantee")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = "Arial"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Claims Data Analysis")
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT
    run.font.name = "Arial"

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Insights from 84 claims across 3 datasets")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = GRAY_TEXT
    run.font.name = "Arial"

    for _ in range(4):
        doc.add_paragraph()

    # Meta info
    for line in [
        "February 2026",
        "Data period: October 2025 \u2013 January 2026",
        "Prepared by ContextBuilder",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_TEXT
        run.font.name = "Arial"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "Arial"
    run.font.small_caps = True

    add_page_break(doc)

    # =========================================================================
    # 1. PORTFOLIO AT A GLANCE
    # =========================================================================
    add_heading1(doc, "1. Portfolio at a Glance")

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Claims analyzed", "84"],
            ["Decision split", "42 approved (50%) / 42 denied (50%)"],
            ["Total approved payout", "CHF 60,885"],
            ["Average approved payout", "CHF 1,450"],
            ["Median approved payout", "CHF 949"],
            ["Vehicle brands represented", "23"],
            ["Top 4 brands (VW, Mercedes, BMW, Audi)", "45 claims (54% of volume)"],
            ["Languages", "DE (43), FR (39), IT (2)"],
            ["Unique garages", "79 across 81 claims with garage data"],
            ["Date range", "8 Oct 2025 \u2013 30 Jan 2026"],
        ],
        col_widths=[3.5, 2.8],
    )

    add_chart(doc, "01_decision_split.png", width=Inches(3.2))

    add_page_break(doc)

    # =========================================================================
    # 2. DECISION PATTERNS
    # =========================================================================
    add_heading1(doc, "2. Decision Patterns \u2014 Why Claims Get Denied")

    add_body(doc,
        'Across the 42 denied claims, a single root cause dominates: "part not covered '
        'by the policy" accounts for 76% of all denials. This is a deterministic, '
        'lookup-based check \u2014 the garage submits a claim for a component, the component '
        "is checked against the policy\u2019s coverage list, and the claim is rejected because "
        "the part is not listed."
    )

    add_heading2(doc, "Denial category breakdown")

    add_table(doc,
        ["Denial Category", "Count", "% of Denials"],
        [
            ["Part not covered", "32", "76%"],
            ["Policy not valid / expired", "4", "10%"],
            ["Consequential / root-cause damage", "4", "10%"],
            ["Wear parts excluded", "1", "2%"],
            ["Mileage exclusion", "1", "2%"],
        ],
        col_widths=[3.2, 1.0, 1.2],
        right_align_cols={1, 2},
    )

    add_chart(doc, "02_denial_reasons.png", width=Inches(5.2))

    add_heading2(doc, "Top denied systems")

    add_body(doc,
        "When the specific parts mentioned in denial reasons are grouped by vehicle system, "
        "the following pattern emerges:"
    )

    add_table(doc,
        ["System", "Denial Mentions"],
        [
            ["Electronics / Software", "10"],
            ["Emissions (EGR, AdBlue, DPF)", "8"],
            ["Seals & Gaskets", "8"],
            ["Engine / Powertrain", "8"],
            ["Cooling / Heating", "5"],
            ["Chassis / Steering", "4"],
            ["Body / Exterior", "2"],
            ["Wear parts", "2"],
            ["Hybrid / EV", "2"],
        ],
        col_widths=[3.5, 1.8],
        right_align_cols={1},
    )

    add_chart(doc, "03_denied_systems.png", width=Inches(5.2))

    add_body(doc,
        "The most frequently denied individual components include software updates, "
        "gaskets/seals, EGR valves, control arms, wiring harnesses, engine control units, "
        "AdBlue systems, timing belts, and water pumps."
    )

    add_callout(doc,
        'The data suggests that a significant portion of denied claims could be identified '
        'earlier in the process, before they require manual adjuster review. The "part not '
        'covered" check is inherently rule-based \u2014 it compares a submitted component against '
        "a defined coverage list, with no judgment call required."
    )

    add_page_break(doc)

    # =========================================================================
    # 3. FINANCIAL PROFILE
    # =========================================================================
    add_heading1(doc, "3. Financial Profile \u2014 Where the Money Goes")

    add_heading2(doc, "Approved payout distribution")

    add_body(doc,
        "All 42 approved claims have payout data. The distribution is right-skewed, "
        "with most claims falling in the CHF 500\u20132,000 range."
    )

    add_table(doc,
        ["Statistic", "Value"],
        [
            ["Mean payout", "CHF 1,450"],
            ["Median payout", "CHF 949"],
            ["Minimum", "CHF 25"],
            ["Maximum", "CHF 4,500"],
            ["25th percentile", "CHF 531"],
            ["75th percentile", "CHF 2,275"],
        ],
        col_widths=[3.0, 2.0],
        right_align_cols={1},
    )

    add_chart(doc, "04_financials.png", width=Inches(5.8))

    add_heading2(doc, "Parts vs. labor split")

    add_body(doc, "Of the 38 approved claims with itemized cost data:")

    add_table(doc,
        ["Component", "Total", "Share"],
        [
            ["Parts", "CHF 30,998", "60%"],
            ["Labor", "CHF 20,695", "40%"],
            ["Combined", "CHF 51,694", "100%"],
        ],
        col_widths=[2.5, 1.8, 1.0],
        right_align_cols={1, 2},
    )

    add_body(doc,
        "Average parts cost per claim: CHF 816. Average labor cost per claim: CHF 545."
    )

    add_heading2(doc, "Deductible structure")

    add_body(doc, "The deductible across 41 approved claims with deductible data:")

    add_bullet(doc, " is the most common deductible, appearing in 20 of 41 claims (49%)", "CHF 150")
    add_bullet(doc, "Mean deductible: CHF 208")
    add_bullet(doc, "Range: CHF 0 \u2013 CHF 500")

    add_heading2(doc, "Reimbursement tiers")

    add_body(doc,
        "23 of 42 approved claims carry a mileage-dependent reimbursement rate that reduces "
        "the payout based on vehicle mileage. The remaining 19 claims either have 100% coverage "
        "or no degradation clause."
    )

    add_table(doc,
        ["Reimbursement Rate", "Claims"],
        [
            ["40%", "10"],
            ["50%", "1"],
            ["60%", "5"],
            ["70%", "3"],
            ["80%", "3"],
            ["90%", "1"],
        ],
        col_widths=[2.5, 1.5],
        right_align_cols={0, 1},
    )

    add_body(doc,
        "The 40% tier is the most common among claims with mileage-dependent reimbursement, "
        "indicating a portfolio with a meaningful share of high-mileage vehicles."
    )

    add_callout(doc,
        "The financial structure of approved claims follows predictable patterns \u2014 "
        "standardized deductibles, formulaic reimbursement rates, and verifiable parts/labor "
        "splits. These characteristics lend themselves well to structured verification."
    )

    add_page_break(doc)

    # =========================================================================
    # 4. BRAND & GARAGE LANDSCAPE
    # =========================================================================
    add_heading1(doc, "4. Brand & Garage Landscape")

    add_heading2(doc, "Brand volume and approval rates")

    add_body(doc,
        "German premium brands dominate the portfolio. The top 4 brands account for "
        "54% of all claims."
    )

    add_table(doc,
        ["Brand", "Claims", "Approved", "Denied", "Approval Rate", "Total Payout (CHF)"],
        [
            ["Volkswagen", "16", "5", "11", "31%", "2,955"],
            ["Mercedes", "14", "8", "6", "57%", "16,831"],
            ["Audi", "8", "5", "3", "63%", "11,048"],
            ["BMW", "7", "5", "2", "71%", "5,264"],
            ["Land Rover", "7", "5", "2", "71%", "7,457"],
            ["Peugeot", "6", "2", "4", "33%", "1,479"],
            ["Ford", "4", "3", "1", "75%", "3,897"],
            ["Mini", "3", "0", "3", "0%", "\u2014"],
            ["Porsche", "2", "0", "2", "0%", "\u2014"],
            ["Jeep", "2", "2", "0", "100%", "2,686"],
            ["Seat", "2", "1", "1", "50%", "863"],
            ["Cupra", "2", "0", "2", "0%", "\u2014"],
        ],
        col_widths=[1.2, 0.7, 0.8, 0.7, 1.0, 1.4],
        right_align_cols={1, 2, 3, 4, 5},
    )

    add_chart(doc, "05_brands.png", width=Inches(5.5))

    add_body(doc, "Remaining 11 brands each have 1 claim.")

    add_body(doc,
        "Volkswagen stands out with the highest volume (16 claims) but the lowest approval "
        'rate among high-volume brands (31%). This is driven primarily by "part not covered" '
        'denials (7 of 11) and "policy expired" denials (4 of 11). No other brand shows the '
        '"policy expired" pattern \u2014 this appears to be Volkswagen-specific in this dataset.'
    )

    add_body(doc,
        "Mercedes generates the highest total payout (CHF 16,831) with 8 approved claims "
        "and an above-average approval rate of 57%."
    )

    add_body(doc,
        "There are no significant brand-specific denial reason patterns beyond "
        'Volkswagen\u2019s expired policies \u2014 "part not covered" is the dominant denial '
        "reason across all brands."
    )

    add_heading2(doc, "Garage landscape")

    add_body(doc,
        "The garage network is overwhelmingly independent: 57 of 81 claims with garage "
        "data come from independent shops. Organized chains and OEM dealerships make up "
        "the remainder."
    )

    add_table(doc,
        ["Garage Type", "Claims", "Approval Rate"],
        [
            ["Independent", "57", "47%"],
            ["AMAG", "7", "86%"],
            ["Emil Frey", "6", "50%"],
            ["Mercedes-Benz (OEM)", "5", "60%"],
            ["BYMYCAR", "3", "33%"],
        ],
        col_widths=[2.5, 1.2, 1.5],
        right_align_cols={1, 2},
    )

    add_body(doc,
        "79 unique garages submitted across 81 claims \u2014 nearly every claim comes from "
        "a different shop. Only 2 garages submitted more than once (Autorama AG and Garage "
        "du Chateau d\u2019en Bas, each with 2 claims)."
    )

    add_body(doc,
        "Claims originate from cities across Switzerland, with the highest concentrations "
        "in Sion (4), Bern (3), and Wetzikon (3)."
    )

    add_page_break(doc)

    # =========================================================================
    # 5. OPERATIONAL OBSERVATIONS
    # =========================================================================
    add_heading1(doc, "5. Operational Observations")

    add_heading2(doc, "Language distribution")

    add_table(doc,
        ["Language", "Claims", "Approved", "Denied", "Approval Rate"],
        [
            ["German (DE)", "43", "21", "22", "49%"],
            ["French (FR)", "39", "21", "18", "54%"],
            ["Italian (IT)", "2", "0", "2", "0%"],
        ],
        col_widths=[1.5, 0.9, 0.9, 0.9, 1.2],
        right_align_cols={1, 2, 3, 4},
    )

    add_body(doc,
        "German and French claims show comparable approval rates (49% vs. 54%) \u2014 there "
        "is no meaningful language bias in decisions. Average approved payouts are nearly "
        "identical: CHF 1,449 (DE) vs. CHF 1,451 (FR)."
    )

    add_body(doc,
        "Italian claims appear for the first time in the eval-v2 dataset (2 claims, both "
        "denied). While too small a sample for conclusions, it may signal emerging activity "
        "in the Italian-speaking market."
    )

    add_body(doc,
        "Bilingual processing (DE/FR) is a baseline requirement for any operational tooling."
    )

    add_heading2(doc, "Document completeness")

    add_body(doc, "Each claim submission averages 4.9 documents. The core document types and their presence rates:")

    add_table(doc,
        ["Document Type", "Present in", "% of Claims"],
        [
            ["Vehicle registration (FZA)", "84 / 84", "100%"],
            ["Policy / guarantee", "84 / 84", "100%"],
            ["Cost estimate (KV)", "83 / 84", "99%"],
            ["Service book", "71 / 84", "85%"],
            ["Mileage proof (KM)", "56 / 84", "67%"],
            ["Photos", "10 / 84", "12%"],
            ["Diagnostic reports", "10 / 84", "12%"],
        ],
        col_widths=[2.5, 1.2, 1.2],
        right_align_cols={2},
    )

    add_body(doc,
        "Cost estimates are present in virtually all claims (99%), and policy documents are "
        "present in 100% of claims. Service books accompany 85% of submissions. Mileage proof "
        "(67%) and diagnostic reports (12%) are less consistently included."
    )

    add_heading2(doc, "Temporal patterns")

    add_table(doc,
        ["Month", "Approved", "Denied", "Total"],
        [
            ["Oct 2025", "7", "0", "7"],
            ["Nov 2025", "4", "0", "4"],
            ["Dec 2025", "14", "17", "31"],
            ["Jan 2026", "17", "25", "42"],
        ],
        col_widths=[1.5, 1.2, 1.0, 1.0],
        right_align_cols={1, 2, 3},
    )

    add_chart(doc, "06_timeline.png", width=Inches(5.0))

    add_body(doc,
        "Volume increases significantly in December 2025 and January 2026. The early months "
        "(Oct\u2013Nov) contain only approved claims \u2014 this reflects the composition of the "
        "eval-v1 dataset, which was enriched first with approved claims. The December\u2013January "
        "spike includes the full decision spectrum and likely reflects both natural volume growth "
        "and the addition of eval-v2 data."
    )

    add_page_break(doc)

    # =========================================================================
    # 6. OPPORTUNITIES THE DATA SUGGESTS
    # =========================================================================
    add_heading1(doc, "6. Opportunities the Data Suggests")

    # -- 1 --
    add_heading2(doc, "1. Early-stage claim screening")

    add_body(doc,
        'The dominance of "part not covered" denials (76% of all denials) suggests that a '
        "pre-check at the point of claim submission \u2014 cross-referencing the claimed part "
        "against the policy\u2019s coverage list \u2014 could surface likely denials before they enter "
        "the manual review queue. This check is deterministic and does not require judgment: "
        "a part is either on the coverage list or it is not."
    )

    # -- 2 --
    add_heading2(doc, "2. Coverage verification support")

    add_body(doc,
        "For approved claims, the financial structure follows highly patterned logic: "
        "standardized deductibles (CHF 150 in nearly half of cases), formulaic mileage-dependent "
        "reimbursement tiers (40%\u201390%), and verifiable parts/labor cost breakdowns. These patterns "
        "lend themselves well to automated verification \u2014 confirming that the math matches the "
        "policy terms."
    )

    # -- 3 --
    add_heading2(doc, "3. Garage guidance")

    add_body(doc,
        "The frequently-denied parts list \u2014 emissions systems (EGR, AdBlue, DPF), electronics "
        "and software, seals and gaskets \u2014 represents an opportunity to provide garages with "
        "clearer upfront visibility into what their customer\u2019s policy covers. Reducing "
        '"part not covered" submissions at the source would decrease processing friction for '
        "both garages and adjusters."
    )

    # -- 4 --
    add_heading2(doc, "4. Product intelligence")

    add_body(doc,
        "The gap between what garages submit and what policies cover \u2014 especially in "
        "emissions systems, electronics, and seals \u2014 reveals where customer expectations "
        "diverge from product coverage. This pattern data is potentially useful for future "
        "product design conversations: understanding which components generate the most claim "
        "friction can inform coverage list evolution."
    )

    add_page_break(doc)

    # =========================================================================
    # 7. DATA NOTES & METHODOLOGY
    # =========================================================================
    add_heading1(doc, "7. Data Notes & Methodology")

    add_heading2(doc, "Datasets")

    add_table(doc,
        ["Dataset", "Claims", "Period", "Notes"],
        [
            ["seed-v1", "4", "Jan 2026", "Initial development set with full ground truth"],
            ["eval-v1", "50", "Oct 2025 \u2013 Jan 2026", "Primary evaluation set"],
            ["eval-v2", "30", "Dec 2025 \u2013 Jan 2026", "Extended set with enriched metadata from provenance PDFs"],
        ],
        col_widths=[1.0, 0.7, 1.5, 3.0],
        right_align_cols={1},
    )

    add_heading2(doc, "Enrichment")

    add_body(doc,
        "The eval-v2 dataset includes additional metadata fields (vehicle details, garage "
        "information, coverage notes, reimbursement rates) extracted from the original claim "
        "PDFs. Earlier datasets have sparser metadata in some fields."
    )

    add_heading2(doc, "Known limitations")

    add_bullet(doc,
        " 84 claims is sufficient for pattern identification but too small for "
        "statistically robust subgroup analysis (e.g., brand-specific denial rates for "
        "low-volume brands)",
        "Sample size:"
    )
    add_bullet(doc,
        " All data comes from NSA Motor Guarantee \u2014 patterns may not "
        "generalize to other warranty products",
        "Single customer:"
    )
    add_bullet(doc,
        " 4 months of data (Oct 2025 \u2013 Jan 2026) \u2014 seasonal effects "
        "cannot be assessed",
        "Short time window:"
    )
    add_bullet(doc,
        " Free-text denial reasons in German and French were "
        "programmatically categorized; automated classification covers ~86% of cases, "
        "with 14% requiring manual review",
        "Denial reason classification:"
    )
    add_bullet(doc,
        " Specific denied parts were extracted via pattern matching on "
        "denial reason text \u2014 some denials reference parts not captured by the extraction "
        "patterns",
        "Part extraction:"
    )

    add_heading2(doc, "Reference")

    add_body(doc,
        "The full interactive analysis with visualizations is available in the project "
        "notebook: analysis/claims_eda.ipynb"
    )

    # =========================================================================
    # FOOTER
    # =========================================================================
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("NSA Motor Guarantee \u2014 Claims Data Analysis  |  Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "Arial"

    # =========================================================================
    # SAVE
    # =========================================================================
    doc.save(str(OUTPUT))
    print(f"Document saved to: {OUTPUT}")


if __name__ == "__main__":
    build()
